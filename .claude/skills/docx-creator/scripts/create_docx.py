#!/usr/bin/env python3
"""
Create a Word document (.docx) from a document JSON spec.

Usage:
  python create_docx.py <document.json> [output.docx]

document.json format:
{
  "metadata": {
    "title": "ドキュメントタイトル",
    "subtitle": "サブタイトル（任意）",
    "author": "著者名",
    "date": "2025-06-13",
    "doc_type": "report | proposal | minutes | manual | seminar",
    "logo": "logo.png",          // assets/ のファイル名（任意）
    "template": "report.docx"    // references/ のテンプレート（任意）
  },
  "style": {
    "primary_color": "1E3A5F",   // 16進数、# なし
    "heading_font": "游明朝",
    "body_font": "游明朝",
    "body_size": 10.5
  },
  "sections": [
    {"type": "title_page"},
    {"type": "toc"},
    {"type": "heading", "text": "1. はじめに", "level": 1},
    {"type": "paragraph", "text": "本文テキスト[1]"},
    {"type": "bullets", "items": ["箇条書き1", "箇条書き2"]},
    {"type": "numbered", "items": ["手順1", "手順2"]},
    {"type": "table", "headers": ["列A", "列B"], "rows": [["a1", "b1"]]},
    {"type": "image", "file": "chart.png", "caption": "図1: 説明"},
    {"type": "page_break"},
    {"type": "references", "citations": [
      {"id": 1, "title": "...", "source": "...", "url": "...", "accessed": "2025-06-13"}
    ]}
  ]
}
"""

import sys
import json
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

SKILL_ROOT = Path(__file__).resolve().parent.parent
ASSETS_DIR = SKILL_ROOT / 'assets'
REFERENCES_DIR = SKILL_ROOT / 'references'


def hex_to_rgb(hex_color: str):
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def resolve_asset(filename: str) -> Path:
    if not filename:
        return None
    p = Path(filename)
    if p.is_absolute() and p.exists():
        return p
    candidate = ASSETS_DIR / filename
    if candidate.exists():
        return candidate
    for f in ASSETS_DIR.rglob(p.name):
        return f
    return None


def resolve_reference(filename: str) -> Path:
    if not filename:
        return None
    candidate = REFERENCES_DIR / filename
    if candidate.exists():
        return candidate
    return None


def set_paragraph_font(para, font_name: str, size_pt: float = None,
                        bold: bool = False, color_hex: str = None, italic: bool = False):
    for run in para.runs:
        run.font.name = font_name
        if size_pt:
            run.font.size = Pt(size_pt)
        run.font.bold = bold
        run.font.italic = italic
        if color_hex:
            run.font.color.rgb = hex_to_rgb(color_hex)


def add_run_with_style(para, text: str, font_name: str, size_pt: float = None,
                        bold: bool = False, color_hex: str = None, italic: bool = False):
    run = para.add_run(text)
    run.font.name = font_name
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = hex_to_rgb(color_hex)
    return run


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def add_page_number(section, doc):
    """フッターにページ番号を追加する。"""
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ---------------------------------------------------------------------------
# Style defaults
# ---------------------------------------------------------------------------

DEFAULT_STYLE = {
    'primary_color':  '1E3A5F',
    'accent_color':   'E8700A',
    'heading_font':   '游明朝',
    'body_font':      '游明朝',
    'title_size':     24.0,
    'h1_size':        16.0,
    'h2_size':        13.0,
    'h3_size':        11.5,
    'body_size':      10.5,
    'caption_size':   9.0,
}

DOC_TYPE_LABELS = {
    'report':   '報告書',
    'proposal': '提案書',
    'minutes':  '議事録',
    'manual':   'マニュアル',
    'seminar':  'セミナー資料',
}


def build_style(override: dict) -> dict:
    style = dict(DEFAULT_STYLE)
    style.update(override)
    return style


# ---------------------------------------------------------------------------
# Section renderers
# ---------------------------------------------------------------------------

def render_title_page(doc: Document, metadata: dict, style: dict):
    """タイトルページを生成する。"""
    doc_type = metadata.get('doc_type', 'report')
    type_label = DOC_TYPE_LABELS.get(doc_type, doc_type)

    # ロゴ
    logo_name = metadata.get('logo', '')
    logo_path = resolve_asset(logo_name)
    if logo_path:
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = para.add_run()
            run.add_picture(str(logo_path), width=Cm(4))
        except Exception as e:
            print(f"  Warning: logo embed failed: {e}")

    # 余白
    for _ in range(4):
        doc.add_paragraph()

    # ドキュメント種別
    type_para = doc.add_paragraph()
    type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_style(type_para, type_label, style['heading_font'],
                        size_pt=12, color_hex=style['accent_color'])

    # タイトル
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_style(title_para, metadata.get('title', ''), style['heading_font'],
                        size_pt=style['title_size'], bold=True, color_hex=style['primary_color'])

    # サブタイトル
    subtitle = metadata.get('subtitle', '')
    if subtitle:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(sub_para, subtitle, style['body_font'],
                            size_pt=14, color_hex=style['primary_color'])

    # 罫線
    for _ in range(3):
        doc.add_paragraph()
    line_para = doc.add_paragraph('─' * 40)
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(2):
        doc.add_paragraph()

    # 著者・日付
    author = metadata.get('author', '')
    doc_date = metadata.get('date', str(date.today()))
    if author:
        a_para = doc.add_paragraph()
        a_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(a_para, author, style['body_font'], size_pt=11)
    d_para = doc.add_paragraph()
    d_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_with_style(d_para, doc_date, style['body_font'], size_pt=11)

    doc.add_page_break()


def render_toc_placeholder(doc: Document, style: dict):
    """目次プレースホルダーを挿入する（Word で更新が必要）。"""
    h = doc.add_heading('目次', level=1)
    for run in h.runs:
        run.font.name = style['heading_font']
        run.font.color.rgb = hex_to_rgb(style['primary_color'])

    para = doc.add_paragraph()
    run = para.add_run('※ Word で開いた後、目次フィールドを右クリック → [フィールドの更新] で目次を生成してください。')
    run.font.name = style['body_font']
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = hex_to_rgb('888888')

    # TOC フィールドコードを挿入
    para2 = doc.add_paragraph()
    run2 = para2.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run2._r.append(fldChar)
    run2._r.append(instrText)
    run2._r.append(fldChar2)

    doc.add_page_break()


def render_heading(doc: Document, text: str, level: int, style: dict):
    h = doc.add_heading(text, level=level)
    sizes = {1: style['h1_size'], 2: style['h2_size'], 3: style['h3_size']}
    for run in h.runs:
        run.font.name = style['heading_font']
        run.font.size = Pt(sizes.get(level, style['h2_size']))
        run.font.color.rgb = hex_to_rgb(style['primary_color'])
    return h


def render_paragraph(doc: Document, text: str, style: dict):
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(style['body_size'])
    add_run_with_style(para, text, style['body_font'], size_pt=style['body_size'])
    return para


def render_bullets(doc: Document, items: list, style: dict, level: int = 0):
    for item in items:
        if isinstance(item, dict):
            # ネスト対応: {"text": "...", "children": [...]}
            para = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
            add_run_with_style(para, item.get('text', ''), style['body_font'],
                                size_pt=style['body_size'])
            if item.get('children'):
                render_bullets(doc, item['children'], style, level + 1)
        else:
            para = doc.add_paragraph(style='List Bullet' if level == 0 else 'List Bullet 2')
            add_run_with_style(para, str(item), style['body_font'], size_pt=style['body_size'])


def render_numbered(doc: Document, items: list, style: dict):
    for item in items:
        para = doc.add_paragraph(style='List Number')
        add_run_with_style(para, str(item), style['body_font'], size_pt=style['body_size'])


def render_table(doc: Document, headers: list, rows: list, style: dict, caption: str = ''):
    if caption:
        cap_para = doc.add_paragraph()
        add_run_with_style(cap_para, caption, style['body_font'],
                            size_pt=style['caption_size'], italic=True)

    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, style['primary_color'])
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(para, str(header), style['heading_font'],
                            size_pt=style['body_size'], bold=True, color_hex='FFFFFF')

    # データ行
    for r_idx, row in enumerate(rows):
        data_row = table.rows[r_idx + 1]
        bg = 'F0F4F8' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_val in enumerate(row):
            cell = data_row.cells[c_idx]
            set_cell_bg(cell, bg)
            add_run_with_style(cell.paragraphs[0], str(cell_val),
                                style['body_font'], size_pt=style['body_size'])

    doc.add_paragraph()  # テーブル後の余白


def render_image(doc: Document, file_name: str, caption: str, style: dict,
                 width_cm: float = 14.0, alt: str = ''):
    img_path = resolve_asset(file_name)
    if not img_path:
        # 視覚的プレースホルダー（表スタイルのボックス）
        table = doc.add_table(rows=3, cols=1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 行1: ラベル
        c0 = table.rows[0].cells[0]
        set_cell_bg(c0, 'EBF3FB')
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(p0, '【 画像プレースホルダー 】',
                            style['heading_font'], size_pt=10, bold=True, color_hex='1E5F9C')

        # 行2: 説明
        c1 = table.rows[1].cells[0]
        set_cell_bg(c1, 'F4F9FF')
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        display_alt = alt if alt else (caption if caption else 'ここに画像が入ります')
        add_run_with_style(p1, display_alt,
                            style['body_font'], size_pt=style['body_size'], color_hex='333333')

        # 行3: ファイルパス指示
        c2 = table.rows[2].cells[0]
        set_cell_bg(c2, 'EBF3FB')
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(p2, f'配置先: assets/{file_name}',
                            style['body_font'], size_pt=style['body_size'] - 1.5,
                            italic=True, color_hex='777777')

        doc.add_paragraph()
        print(f"  [PLACEHOLDER] {file_name}: {display_alt[:50]}")
        return

    try:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(str(img_path), width=Cm(width_cm))
    except Exception as e:
        print(f"  Warning: image embed failed ({file_name}): {e}")
        return

    if caption:
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_with_style(cap_para, caption, style['body_font'],
                            size_pt=style['caption_size'], italic=True,
                            color_hex='555555')

    doc.add_paragraph()


def render_references(doc: Document, citations: list, style: dict):
    """末尾に番号付き参考文献リストを生成する。"""
    doc.add_page_break()
    render_heading(doc, '参考文献', level=1, style=style)

    for cite in citations:
        cid = cite.get('id', '')
        title = cite.get('title', '')
        source = cite.get('source', '')
        url = cite.get('url', '')
        accessed = cite.get('accessed', '')
        note = cite.get('note', '')

        parts = [f'[{cid}]  {title}']
        if source:
            parts.append(f'  {source}')
        if url:
            parts.append(f'  {url}')
        if accessed:
            parts.append(f'  （アクセス日: {accessed}）')
        if note:
            parts.append(f'  ※ {note}')

        para = doc.add_paragraph()
        add_run_with_style(para, '\n'.join(parts), style['body_font'],
                            size_pt=style['body_size'] - 0.5)
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_after = Pt(6)


# ---------------------------------------------------------------------------
# Minutes / Seminar helpers
# ---------------------------------------------------------------------------

def render_minutes_header(doc: Document, metadata: dict, style: dict):
    """議事録用のヘッダー情報テーブルを生成する。"""
    fields = [
        ('会議名', metadata.get('title', '')),
        ('日時',   metadata.get('date', '')),
        ('場所',   metadata.get('location', '')),
        ('出席者', metadata.get('attendees', '')),
        ('記録者', metadata.get('author', '')),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate(fields):
        cell_l = table.rows[i].cells[0]
        set_cell_bg(cell_l, 'EEF2F6')
        add_run_with_style(cell_l.paragraphs[0], label, style['heading_font'],
                            size_pt=style['body_size'], bold=True)
        add_run_with_style(table.rows[i].cells[1].paragraphs[0], str(value),
                            style['body_font'], size_pt=style['body_size'])
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

SECTION_RENDERERS = {
    'title_page': lambda doc, sec, meta, style:
        render_title_page(doc, meta, style),
    'toc': lambda doc, sec, meta, style:
        render_toc_placeholder(doc, style),
    'minutes_header': lambda doc, sec, meta, style:
        render_minutes_header(doc, meta, style),
    'heading': lambda doc, sec, meta, style:
        render_heading(doc, sec.get('text', ''), sec.get('level', 1), style),
    'paragraph': lambda doc, sec, meta, style:
        render_paragraph(doc, sec.get('text', ''), style),
    'bullets': lambda doc, sec, meta, style:
        render_bullets(doc, sec.get('items', []), style),
    'numbered': lambda doc, sec, meta, style:
        render_numbered(doc, sec.get('items', []), style),
    'table': lambda doc, sec, meta, style:
        render_table(doc, sec.get('headers', []), sec.get('rows', []), style,
                     caption=sec.get('caption', '')),
    'image': lambda doc, sec, meta, style:
        render_image(doc, sec.get('file', ''), sec.get('caption', ''), style,
                     width_cm=sec.get('width_cm', 14.0), alt=sec.get('alt', '')),
    'page_break': lambda doc, sec, meta, style:
        doc.add_page_break(),
    'references': lambda doc, sec, meta, style:
        render_references(doc, sec.get('citations', []), style),
}


def main():
    if len(sys.argv) < 2:
        print("Usage: create_docx.py <document.json> [output.docx]")
        sys.exit(1)

    doc_path = Path(sys.argv[1])
    with open(doc_path, 'r', encoding='utf-8') as f:
        spec = json.load(f)

    metadata = spec.get('metadata', {})
    style = build_style(spec.get('style', {}))
    sections = spec.get('sections', [])

    # テンプレートがあれば使う
    template_name = metadata.get('template', '')
    template_path = resolve_reference(template_name) if template_name else None
    if template_path:
        doc = Document(str(template_path))
        print(f"Template: {template_path.name}")
    else:
        doc = Document()
        # ページ設定（A4）
        section = doc.sections[0]
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        add_page_number(section, doc)

    # アセット一覧を表示
    if ASSETS_DIR.exists():
        asset_files = [f.name for f in ASSETS_DIR.iterdir() if not f.name.startswith('.')]
        if asset_files:
            print(f"Assets available: {asset_files}")

    # セクションをレンダリング
    for sec in sections:
        sec_type = sec.get('type', 'paragraph')
        renderer = SECTION_RENDERERS.get(sec_type)
        if renderer:
            renderer(doc, sec, metadata, style)
        else:
            print(f"  Unknown section type: {sec_type}")

    output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else doc_path.with_suffix('.docx')
    doc.save(str(output_path))
    print(f"Saved: {output_path}  ({len(sections)} sections)")

    # 最終実行コマンドを保存（画像差し替え時の自動再実行用）
    last_run = {
        "document_json": str(doc_path.absolute()),
        "output_docx": str(output_path.absolute()),
        "rerun_cmd": f"python3 {Path(__file__).absolute()} {doc_path.absolute()} {output_path.absolute()}"
    }
    last_run_path = Path(__file__).resolve().parent.parent / 'last_run.json'
    with open(last_run_path, 'w', encoding='utf-8') as f:
        json.dump(last_run, f, ensure_ascii=False, indent=2)
    print(f"RERUN_CMD: {last_run['rerun_cmd']}")


if __name__ == '__main__':
    main()
